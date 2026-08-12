from __future__ import annotations
from pathlib import Path
from docx import Document
from trump_monitor.models import RunResult

def export_docx(result: RunResult,path: str|Path):
    out=Path(path); out.parent.mkdir(parents=True,exist_ok=True); d=Document()
    d.add_heading('川普72小時事件與市場影響報告｜English + 繁體中文',0)
    d.add_paragraph(f'Run ID：{result.run_id}｜狀態：{result.status}｜Truth：{result.truth_social_status}')
    for e in result.events:
        d.add_heading('★'*e.score.importance+' '+e.topic,1)
        d.add_paragraph('中文：'+(e.topic_zh or '翻譯未取得'))
        d.add_paragraph('English summary：'+e.summary)
        d.add_paragraph('中文摘要：'+(e.summary_zh or '翻譯未取得'))
        d.add_paragraph(f'類別：{e.category}｜可信度：{e.score.confidence:.0%}｜重大性：{e.materiality_score}/100｜GTC：{e.battle_action}')
        d.add_paragraph('受惠：'+'、'.join(e.beneficiary_sectors)); d.add_paragraph('受壓：'+'、'.join(e.negative_sectors))
        d.add_paragraph('翻譯狀態：'+e.translation_status)
    d.save(out); return out
